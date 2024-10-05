# File: gui/base_projects_tab.py

from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QTreeWidget, QTreeWidgetItem,
    QHBoxLayout, QMessageBox, QCheckBox, QPushButton, QFileDialog, QMenu, QAction
)
from PyQt5.QtGui import QColor
from PyQt5.QtCore import Qt, QPoint
from datetime import datetime
import os
import sys
import shutil
import tempfile

from logger import get_logger
from utils import (
    sanitize_filename, get_project_dir, get_template_dir, get_project_folder_name, open_docx_file
)
from gui.widgets.buttons import SplitButton
from gui.event_handlers import (
    handle_project_delete, handle_toggle_unit_status, handle_move_project,
    handle_import_floor_plan, handle_import_master_floor_plan
)
from controllers.project_controller import ProjectController
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from database import UnitModel  # Removed ProjectModel import since we're using the controller

logger = get_logger(__name__)

class BaseProjectsTab(QWidget):
    def __init__(self, db, status_filter=None, title=""):
        super().__init__()
        self.db = db
        self.controller = ProjectController(self.db)
        self.status_filter = status_filter
        self.title = title
        self.template_dir = get_template_dir()
        self.layout = QVBoxLayout()
        self.setLayout(self.layout)
        self.setup_ui()
        self.load_projects()

    def setup_ui(self):
        # Buttons Layout
        self.buttons_layout = QHBoxLayout()
        
        # View DOCX Split Button using SplitButton class
        self.view_docx_split_btn = SplitButton(
            text="View DOCX",
            tooltip=f"View {self.title} DOCX",
            default_action=self.view_docx_overview,
            menu_actions=[
                ("Save As...", self.save_docx_overview)
            ]
        )
        self.buttons_layout.addWidget(self.view_docx_split_btn)

        self.layout.addLayout(self.buttons_layout)

        # Projects Tree
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels([
            "Project Name",
            "Project Number",
            "Main Contractor",
            "Completed Units",
            "Status",
            "Extra",
            "Innregulering",
            "Sjekkliste",
            "Floor Plan(s)",
            "Move(1)",
            "Move(2)",
            "Start Date",
            "End Date",
            "Worker"
        ])
        self.tree.setColumnCount(14)
        self.tree.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self.open_context_menu)
        self.layout.addWidget(self.tree)

        # Temporary DOCX path
        self.docx_path = os.path.join(tempfile.gettempdir(), f"{sanitize_filename(self.title)}_Projects.docx")

    def load_projects(self):
        self.tree.clear()
        projects = self.controller.load_projects(status=self.status_filter)
        logger.info(f"Loading {len(projects)} projects into the '{self.title}' tab.")
        for project in projects:
            project_item = QTreeWidgetItem([
                project.name,
                project.number,
                project.main_contractor if project.main_contractor else "",
                "",
                project.status,
                project.extra if project.extra else "",
                "",
                "",
                "",
                "",
                "",
                self.format_date(project.start_date),
                self.format_date(project.end_date) if project.end_date else "",
                project.worker
            ])
            project_item.setData(0, Qt.UserRole, project.id)
            self.tree.addTopLevelItem(project_item)
            if project.is_residential_complex and project.units:
                # Calculate completed units
                completed = sum(1 for unit in project.units if self.get_unit_status(project.id, unit))
                total = len(project.units)
                project_item.setText(3, f"{completed}/{total}")

                for unit_name in project.units:
                    unit_item = QTreeWidgetItem([
                        unit_name,
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        ""
                    ])
                    # Add a checkbox for done/undone
                    checkbox = QCheckBox()
                    # Retrieve the unit's current status
                    is_done = self.get_unit_status(project.id, unit_name)
                    checkbox.setChecked(is_done)
                    # Connect the checkbox state change to the event handler
                    checkbox.stateChanged.connect(
                        lambda state, p=project, u=unit_name: handle_toggle_unit_status(self.controller, p, u, state, self)
                    )
                    self.tree.setItemWidget(unit_item, 0, checkbox)
                    project_item.addChild(unit_item)

                    # Innregulering Split Button for Unit
                    innregulering_split_btn = SplitButton(
                        text="View",
                        tooltip="View Innregulering DOCX",
                        default_action=lambda checked=False, p=project, u=unit_name: self.view_docx(p, "Innregulering", u),
                        menu_actions=[
                            ("Save As...", lambda checked=False, p=project, u=unit_name: self.save_docx_as(p, "Innregulering", u))
                        ]
                    )
                    self.tree.setItemWidget(unit_item, 6, innregulering_split_btn)

                    # Sjekkliste Split Button for Unit
                    sjekkliste_split_btn = SplitButton(
                        text="View",
                        tooltip="View Sjekkliste DOCX",
                        default_action=lambda checked=False, p=project, u=unit_name: self.view_docx(p, "Sjekkliste", u),
                        menu_actions=[
                            ("Save As...", lambda checked=False, p=project, u=unit_name: self.save_docx_as(p, "Sjekkliste", u))
                        ]
                    )
                    self.tree.setItemWidget(unit_item, 7, sjekkliste_split_btn)

                    # Floor Plan(s) Split Button for Unit
                    floor_plan_split_btn = SplitButton(
                        text="View",
                        tooltip="View Floor Plan(s)",
                        default_action=lambda checked=False, p=project, u=unit_name: self.view_floor_plan(p, u),
                        menu_actions=[
                            ("Import PDF", lambda checked=False, p=project, u=unit_name: handle_import_floor_plan(p, u, self)),
                            ("Save As...", lambda checked=False, p=project, u=unit_name: self.save_floor_plan_as(p, u))
                        ]
                    )
                    self.tree.setItemWidget(unit_item, 8, floor_plan_split_btn)

                # Add Master Split Button for Residential Projects in "Floor Plan(s)" column
                master_split_btn = SplitButton(
                    text="Master",
                    tooltip="Manage Master Floor Plan",
                    default_action=lambda checked=False, p=project: self.view_master_floor_plan(p),
                    menu_actions=[
                        ("Import PDF", lambda checked=False, p=project: handle_import_master_floor_plan(p, self)),
                        ("Save As...", lambda checked=False, p=project: self.save_master_floor_plan_as(p))
                    ]
                )
                self.tree.setItemWidget(project_item, 8, master_split_btn)

                # Expand the project item to show unit names by default
                self.tree.expandItem(project_item)
            else:
                # Innregulering Split Button
                innregulering_split_btn = SplitButton(
                    text="View",
                    tooltip="View Innregulering DOCX",
                    default_action=lambda checked=False, p=project: self.view_docx(p, "Innregulering"),
                    menu_actions=[
                        ("Save As...", lambda checked=False, p=project: self.save_docx_as(p, "Innregulering"))
                    ]
                )
                self.tree.setItemWidget(project_item, 6, innregulering_split_btn)

                # Sjekkliste Split Button
                sjekkliste_split_btn = SplitButton(
                    text="View",
                    tooltip="View Sjekkliste DOCX",
                    default_action=lambda checked=False, p=project: self.view_docx(p, "Sjekkliste"),
                    menu_actions=[
                        ("Save As...", lambda checked=False, p=project: self.save_docx_as(p, "Sjekkliste"))
                    ]
                )
                self.tree.setItemWidget(project_item, 7, sjekkliste_split_btn)

                # Floor Plan(s) Split Button for Project
                floor_plan_split_btn = SplitButton(
                    text="View",
                    tooltip="View Floor Plan(s)",
                    default_action=lambda checked=False, p=project: self.view_floor_plan(p),
                    menu_actions=[
                        ("Import PDF", lambda checked=False, p=project: handle_import_floor_plan(p, None, self)),
                        ("Save As...", lambda checked=False, p=project: self.save_floor_plan_as(p))
                    ]
                )
                self.tree.setItemWidget(project_item, 8, floor_plan_split_btn)

                # Move(1) Button
                move1_btn = QPushButton("Active")
                move1_btn.setToolTip("Move Project to Active")
                move1_btn.setStyleSheet("background-color: yellow")
                move1_btn.clicked.connect(lambda checked=False, p=project: handle_move_project(self.controller, p, "Active", self))
                self.tree.setItemWidget(project_item, 9, move1_btn)

                # Move(2) Button
                move2_btn = QPushButton("Completed")
                move2_btn.setToolTip("Move Project to Completed")
                move2_btn.setStyleSheet("background-color: green")
                move2_btn.clicked.connect(lambda checked=False, p=project: handle_move_project(self.controller, p, "Completed", self))
                self.tree.setItemWidget(project_item, 10, move2_btn)

                if not project.is_residential_complex:
                    # Set Completed Units to N/A
                    project_item.setText(3, "N/A")

                logger.debug(f"Added project '{project.name}' with status '{project.status}' to the tree.")

    def get_unit_status(self, project_id, unit_name):
        """
        Helper method to get the status of a unit.
        """
        unit = self.controller.db.session.query(UnitModel).filter_by(project_id=project_id, name=unit_name).first()
        return unit.is_done if unit else False

    def open_context_menu(self, position: QPoint):
        item = self.tree.itemAt(position)
        if item and not item.parent():
            menu = QMenu(self)
            edit_action = QAction("Edit", self)
            delete_action = QAction("Delete", self)
            menu.addAction(edit_action)
            menu.addAction(delete_action)
            action = menu.exec_(self.tree.viewport().mapToGlobal(position))
            if action == edit_action:
                # Placeholder for Edit functionality
                pass
            elif action == delete_action:
                project_id = item.data(0, Qt.UserRole)
                handle_project_delete(self.controller, project_id, self)

    def view_docx_overview(self):
        self.generate_docx()
        if not os.path.exists(self.docx_path):
            QMessageBox.warning(self, "DOCX Error", f"{self.title} DOCX does not exist.")
            logger.warning(f"{self.title} DOCX not found.")
            return

        success, message = open_docx_file(self.docx_path)
        if not success:
            QMessageBox.warning(self, "DOCX Error", message)
            logger.error(f"Failed to open {self.title} DOCX: {message}")

    def save_docx_overview(self):
        self.generate_docx()
        options = QFileDialog.Options()
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Overview As...",
            f"{sanitize_filename(self.title)}_Projects.docx",
            "Word Documents (*.docx)",
            options=options
        )
        if save_path:
            try:
                shutil.copy(self.docx_path, save_path)
                QMessageBox.information(self, "Success", f"Overview saved successfully at:\n{save_path}")
                logger.info(f"Overview saved as {save_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save Overview:\n{str(e)}")
                logger.error(f"Failed to save Overview DOCX: {e}")

    def generate_docx(self):
        try:
            document = Document()
            # Set page orientation to landscape
            section = document.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

            # Add header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = self.title
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add table
            projects = self.controller.load_projects(status=self.status_filter)
            if not projects:
                document.add_paragraph("No projects to display.")
            else:
                table = document.add_table(rows=1, cols=8)
                table.style = 'Light List Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Project Name'
                hdr_cells[1].text = 'Project Number'
                hdr_cells[2].text = 'Main Contractor'
                hdr_cells[3].text = 'Completed Units'
                hdr_cells[4].text = 'Status'
                hdr_cells[5].text = 'Start Date'
                hdr_cells[6].text = 'End Date'
                hdr_cells[7].text = 'Worker'

                for project in projects:
                    row_cells = table.add_row().cells
                    row_cells[0].text = project.name
                    row_cells[1].text = project.number
                    row_cells[2].text = project.main_contractor if project.main_contractor else "N/A"
                    if project.is_residential_complex:
                        completed = sum(1 for unit in project.units if self.get_unit_status(project.id, unit))
                        total = len(project.units)
                        row_cells[3].text = f"{completed}/{total}"
                    else:
                        row_cells[3].text = "N/A"
                    row_cells[4].text = project.status
                    row_cells[5].text = self.format_date(project.start_date)
                    row_cells[6].text = self.format_date(project.end_date) if project.end_date else "N/A"
                    row_cells[7].text = project.worker

                    # Adjust column widths to fit the page
                    widths = [Inches(1.5), Inches(1.0), Inches(1.5), Inches(1.2), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.2)]
                    for row in table.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width

            # Add footer with current date
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = datetime.now().strftime("%d-%m-%Y")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            document.save(self.docx_path)
            logger.info(f"Generated DOCX at {self.docx_path}")
        except Exception as e:
            QMessageBox.critical(self, "DOCX Generation Error", f"Failed to generate DOCX:\n{str(e)}")
            logger.error(f"Failed to generate DOCX at {self.docx_path}: {e}")

    def save_docx_as(self, project, doc_type, unit_name=None):
        options = QFileDialog.Options()
        if unit_name:
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                f"Save {doc_type} As for {unit_name}...",
                f"{sanitize_filename(project.name)}_{sanitize_filename(unit_name)}_{doc_type}.docx",
                "Word Documents (*.docx)",
                options=options
            )
        else:
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                f"Save {doc_type} As...",
                f"{sanitize_filename(project.name)}_{doc_type}.docx",
                "Word Documents (*.docx)",
                options=options
            )
        if save_path:
            try:
                folder_name = get_project_folder_name(project)
                if unit_name:
                    docx_file = os.path.join(get_project_dir(), folder_name, sanitize_filename(unit_name), f"{doc_type}.docx")
                else:
                    docx_file = os.path.join(get_project_dir(), folder_name, f"{doc_type}.docx")
                if not os.path.exists(docx_file):
                    QMessageBox.warning(self, "File Not Found", f"{doc_type}.docx does not exist.")
                    logger.warning(f"{doc_type}.docx not found for project '{project.name}'" + (f" and unit '{unit_name}'." if unit_name else "."))
                    return
                shutil.copy(docx_file, save_path)
                QMessageBox.information(self, "Success", f"{doc_type} saved successfully at:\n{save_path}")
                logger.info(f"{doc_type} saved as {save_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save {doc_type}:\n{str(e)}")
                logger.error(f"Failed to save {doc_type} for project '{project.name}'" + (f" and unit '{unit_name}': {e}" if unit_name else f": {e}"))

    def view_docx(self, project, doc_type, unit_name=None):
        if unit_name:
            folder_name = get_project_folder_name(project)
            docx_file = os.path.join(get_project_dir(), folder_name, sanitize_filename(unit_name), f"{doc_type}.docx")
        else:
            folder_name = get_project_folder_name(project)
            docx_file = os.path.join(get_project_dir(), folder_name, f"{doc_type}.docx")

        if not os.path.exists(docx_file):
            QMessageBox.warning(self, "DOCX Error", f"{doc_type}.docx does not exist for this {'unit ' + unit_name if unit_name else 'project'}.")
            logger.warning(f"{doc_type}.docx not found for project '{project.name}'" + (f" and unit '{unit_name}'." if unit_name else "."))
            return

        success, message = open_docx_file(docx_file)
        if not success:
            QMessageBox.warning(self, "DOCX Error", message)
            logger.error(f"Failed to open {doc_type}.docx for project '{project.name}'" + (f" and unit '{unit_name}': {message}" if unit_name else f": {message}"))

    def view_floor_plan(self, project, unit_name=None):
        if project.is_residential_complex and not unit_name:
            QMessageBox.warning(self, "Floor Plan Error", "This project is a residential complex. Please select a unit to view its floor plan.")
            logger.warning(f"Attempted to view project-level floor plan for residential project '{project.name}'.")
            return

        try:
            folder_name = get_project_folder_name(project)

            if unit_name:
                floor_plan_file = os.path.join(
                    get_project_dir(),
                    folder_name,
                    sanitize_filename(unit_name),
                    "Floor plan",
                    "FloorPlan.pdf"
                )
            else:
                floor_plan_file = os.path.join(
                    get_project_dir(),
                    folder_name,
                    "Floor plan",
                    "FloorPlan.pdf"
                )

            if not os.path.exists(floor_plan_file):
                QMessageBox.warning(self, "Floor Plan Error", "Floor Plan PDF does not exist.")
                logger.warning(f"Floor Plan PDF not found for project '{project.name}'" + (f" and unit '{unit_name}'." if unit_name else "."))
                return

            try:
                if sys.platform == "win32":
                    os.startfile(floor_plan_file)
                elif sys.platform == "darwin":
                    subprocess.call(["open", floor_plan_file])
                else:
                    subprocess.call(["xdg-open", floor_plan_file])
                logger.info(f"Opened Floor Plan PDF: {floor_plan_file}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to open Floor Plan PDF:\n{str(e)}")
                logger.error(f"Failed to open Floor Plan PDF '{floor_plan_file}': {e}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred while trying to view the Floor Plan:\n{str(e)}")
            logger.error(f"Error in view_floor_plan for project '{project.name}': {e}")

    def save_floor_plan_as(self, project, unit_name=None):
        options = QFileDialog.Options()
        if unit_name:
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                f"Save Floor Plan As for {unit_name}...",
                f"{sanitize_filename(project.name)}_{sanitize_filename(unit_name)}_FloorPlan.pdf",
                "PDF Files (*.pdf)",
                options=options
            )
        else:
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Floor Plan As...",
                f"{sanitize_filename(project.name)}_FloorPlan.pdf",
                "PDF Files (*.pdf)",
                options=options
            )
        if save_path:
            try:
                folder_name = get_project_folder_name(project)
                if unit_name:
                    floor_plan_file = os.path.join(get_project_dir(), folder_name, sanitize_filename(unit_name), "Floor plan", "FloorPlan.pdf")
                else:
                    floor_plan_file = os.path.join(get_project_dir(), folder_name, "Floor plan", "FloorPlan.pdf")
                if not os.path.exists(floor_plan_file):
                    QMessageBox.warning(self, "File Not Found", "Floor Plan PDF does not exist.")
                    logger.warning(f"Floor Plan PDF not found for project '{project.name}'" + (f" and unit '{unit_name}'." if unit_name else "."))
                    return
                shutil.copy(floor_plan_file, save_path)
                QMessageBox.information(self, "Success", f"Floor Plan saved successfully at:\n{save_path}")
                logger.info(f"Floor Plan saved as {save_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save Floor Plan:\n{str(e)}")
                logger.error(f"Failed to save Floor Plan for project '{project.name}'" + (f" and unit '{unit_name}': {e}" if unit_name else f": {e}"))

    def view_master_floor_plan(self, project):
        try:
            folder_name = get_project_folder_name(project)
            master_floor_plan_file = os.path.join(get_project_dir(), folder_name, "Master", "MasterFloorPlan.pdf")

            if not os.path.exists(master_floor_plan_file):
                QMessageBox.warning(self, "Master Floor Plan Error", "Master Floor Plan PDF does not exist.")
                logger.warning(f"Master Floor Plan PDF not found for project '{project.name}'.")
                return

            try:
                if sys.platform == "win32":
                    os.startfile(master_floor_plan_file)
                elif sys.platform == "darwin":
                    subprocess.call(["open", master_floor_plan_file])
                else:
                    subprocess.call(["xdg-open", master_floor_plan_file])
                logger.info(f"Opened Master Floor Plan PDF: {master_floor_plan_file}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to open Master Floor Plan PDF:\n{str(e)}")
                logger.error(f"Failed to open Master Floor Plan PDF '{master_floor_plan_file}': {e}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred while trying to view the Master Floor Plan:\n{str(e)}")
            logger.error(f"Error in view_master_floor_plan for project '{project.name}': {e}")

    def save_master_floor_plan_as(self, project):
        options = QFileDialog.Options()
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Master Floor Plan As...",
            f"{sanitize_filename(project.name)}_MasterFloorPlan.pdf",
            "PDF Files (*.pdf)",
            options=options
        )
        if save_path:
            try:
                folder_name = get_project_folder_name(project)
                master_floor_plan_file = os.path.join(get_project_dir(), folder_name, "Master", "MasterFloorPlan.pdf")
                if not os.path.exists(master_floor_plan_file):
                    QMessageBox.warning(self, "File Not Found", "Master Floor Plan PDF does not exist.")
                    logger.warning(f"Master Floor Plan PDF not found for project '{project.name}'.")
                    return
                shutil.copy(master_floor_plan_file, save_path)
                QMessageBox.information(self, "Success", f"Master Floor Plan saved successfully at:\n{save_path}")
                logger.info(f"Master Floor Plan saved as {save_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save Master Floor Plan:\n{str(e)}")
                logger.error(f"Failed to save Master Floor Plan for project '{project.name}': {e}")

    def format_date(self, date_str):
        try:
            return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d-%m-%Y")
        except:
            return date_str
